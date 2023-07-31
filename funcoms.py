import pandas as pd
import openpyxl
from docx import Document
from langchain.document_loaders import JSONLoader
import jq
import json
from pathlib import Path

class JSONMod(JSONLoader):
    def __init__(self, file_path):
        file_path = Path(file_path)
        super().__init__(file_path, jq_schema='.', text_content=False)
        self.file_path = file_path
        self.metadata = None

    def load_metadata(self):
        with open(self.file_path, 'r') as f:
            json_data = json.load(f)
        metadata = jq.compile('.').input(json_data).all()
        return metadata

    def load(self):
        self.metadata = self.load_metadata()
        return super().load()


class WordLoader:
    def __init__(self, filepath):
        self.filepath = filepath

    def load(self):
        doc =docx. Document(self.filepath)
        for paragraph in doc.paragraphs:
            yield Document(text=paragraph.text)

import os
import pandas as pd

def get_loader_cls(file_extension: str):
    from langchain.document_loaders import TextLoader, PDFMinerLoader, BSHTMLLoader, JSONLoader  # import any other loaders you might
    
    if file_extension == '.txt':
        return TextLoader
    elif file_extension == '.pdf':
        return PDFMinerLoader
    elif file_extension == '.html':
        return BSHTMLLoader
    elif file_extension == '.docx':
        return WordLoader
    elif file_extension == '.json':
        return JSONMod
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")
